from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
REPORT_DIR = BASE_DIR / "reports"
OUTPUT_PATH = REPORT_DIR / "中文指代消解与句子改写系统_最终报告.docx"


def set_cell_shading(cell, fill: str) -> None:
    """Apply background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    """Set deterministic table cell width."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    """Set fixed table width and left indent."""
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    """Set table cell margins in DXA."""
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tbl_cell_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("中文指代消解与句子改写系统项目报告")
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run("基于规则增强、错误分析与 Streamlit 可视化的中文 NLP 期末项目")
    sub.font.size = Pt(11)
    sub.font.color.rgb = RGBColor(71, 85, 105)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("项目仓库：github.com/gingko123/chinese-coreference-rewrite")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(71, 85, 105)

    meta = doc.add_table(rows=7, cols=2)
    meta.style = "Table Grid"
    meta.autofit = False
    set_table_width(meta, width_dxa=7200, indent_dxa=1080)
    set_cell_margins(meta)
    rows = [
        ("学生姓名", "郭英恺、刘毅恒"),
        ("学号", "10245102433、10245102406"),
        ("班级", "________________"),
        ("课程项目", "自然语言处理期末项目"),
        ("项目主题", "中文指代消解与句子改写"),
        ("实现形式", "Python + Streamlit + 规则增强 + 数据集评估"),
        ("报告日期", "2026 年 6 月"),
    ]
    for row, (key, value) in zip(meta.rows, rows):
        set_cell_width(row.cells[0], 2200)
        set_cell_width(row.cells[1], 5000)
        set_cell_shading(row.cells[0], "F2F4F7")
        row.cells[0].paragraphs[0].add_run(key).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    doc.add_paragraph()


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(31, 41, 55)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)
    set_cell_margins(table)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, "F2F4F7")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 and len(value) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            run.font.size = Pt(10)

    doc.add_paragraph()


def load_json(path: str) -> list[dict]:
    return json.loads((BASE_DIR / path).read_text(encoding="utf-8"))


def build_report() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("摘要", level=1)
    add_para(
        doc,
        "本项目实现了一个面向中文文本的指代消解与句子改写系统。系统能够在用户输入中文句子后，"
        "识别候选实体与代词，预测代词对应的先行词，并生成指代更加明确的改写文本。项目采用"
        "规则增强路线，强调可解释性和可展示性，同时提供 Streamlit Web Demo、数据集评估、"
        "消融实验和错误分析功能。新增 85 条真实风格标注数据后，系统自动生成 train.json 与"
        "rewrite 字段，并通过组织指代规则增强将训练集 F1 从 0.7963 提升到 0.8765。"
    )
    add_para(
        doc,
        "报告围绕任务背景、系统设计、核心规则、数据构建、实验结果、错误分析和部署展示展开。"
        "与单纯调用现成模型不同，本项目更强调从数据标注、规则设计到可视化验证的完整工程流程，"
        "适合作为小规模中文 NLP 系统设计与实验分析案例。"
    )

    doc.add_heading("1. 项目背景与意义", level=1)
    add_para(
        doc,
        "指代消解是自然语言处理中的重要任务，目标是判断文本中的代词或指代表达所对应的先行实体。"
        "在中文语境中，代词、省略和上下文依赖现象非常常见，例如“他”“她”“它”“该公司”“这家车企”等表达"
        "往往需要结合前文实体才能理解。若系统无法正确识别这些关系，后续的问答、摘要、信息抽取和文本改写都会受到影响。"
    )
    add_para(
        doc,
        "本项目将任务定位为“中文指代消解 + 句子改写”。相比只输出标签，句子改写能更直观地展示系统理解结果，"
        "例如将“马斯克整理了书，他准备再次确认它”改写为“马斯克整理了书，马斯克准备再次确认书”。"
    )
    doc.add_heading("1.1 研究问题", level=2)
    add_para(
        doc,
        "本项目关注三个具体问题：第一，如何在不依赖大型模型的情况下识别中文文本中的实体和代词；"
        "第二，如何用可解释规则判断代词与先行词之间的关系；第三，如何通过可视化和错误分析展示系统效果，"
        "并找出后续引入语义模型的方向。"
    )

    doc.add_heading("2. 项目目标", level=1)
    add_bullets(doc, [
        "支持用户输入中文文本，并自动识别实体、代词和指代表达。",
        "为每个代词生成候选先行词，并根据规则分数选择最可能的指代对象。",
        "输出可解释的指代关系，包括候选实体、分数和规则理由。",
        "根据指代结果生成更明确的改写文本。",
        "提供 Streamlit 可视化页面、数据集评估、消融实验和错误分析。",
        "支持新标注数据接入，自动生成训练集与错误分析报告。",
    ])
    doc.add_heading("2.1 项目创新点", level=2)
    add_bullets(doc, [
        "将指代消解结果进一步转化为句子改写结果，使系统输出更直观。",
        "保留候选实体得分和规则理由，避免黑盒式展示。",
        "新增训练数据自动接入流程，可从未 rewrite 的标注集生成 train.json。",
        "加入错误类型归因报告，把实验结果转化为可写入报告的分析结论。",
        "实现本地运行和云服务器展示两种使用方式，便于课堂演示。",
    ])
    doc.add_heading("2.2 使用场景与目标用户", level=2)
    add_table(
        doc,
        ["使用场景", "目标用户", "实际价值"],
        [
            ["课堂演示", "自然语言处理课程学生与教师", "通过可视化页面观察指代消解过程，便于理解 NLP 任务流程。"],
            ["文本改写辅助", "需要处理中文材料的写作者", "将含糊代词替换为明确实体，使句子更适合阅读和二次加工。"],
            ["问答/摘要预处理", "中文 NLP 应用开发者", "在进入问答、摘要或信息抽取前，先显式化上下文中的指代关系。"],
            ["数据标注与错误分析", "小规模 NLP 项目开发者", "快速构建样本、检查错误类型，并决定后续规则或模型改进方向。"],
        ],
        [1900, 2500, 4960],
    )

    doc.add_heading("3. 技术栈与项目结构", level=1)
    add_table(
        doc,
        ["模块", "技术或文件", "主要作用"],
        [
            ["前端展示", "Streamlit / app.py", "实现文本输入、结果展示、数据集评估和消融实验页面。"],
            ["实体与代词抽取", "src/mention_extractor.py", "识别人名、物体、组织及代词，并接入训练词表。"],
            ["指代消解", "src/resolver.py", "根据距离、位置、类型、性别和角色模式对候选先行词打分。"],
            ["文本改写", "src/rewriter.py", "用先行词替换可消解代词，歧义代词跳过自动改写。"],
            ["实验评估", "src/evaluator.py", "计算 Accuracy、Precision、Recall、F1 并生成错误样本。"],
            ["增强接口", "HanLP / LTP", "预留中文 NLP 后端接口，未安装时自动回退规则版。"],
        ],
        [1800, 2400, 5160],
    )

    doc.add_heading("4. 系统设计", level=1)
    add_para(doc, "系统整体流程如下：")
    add_code_block(
        doc,
        "输入文本 -> Mention 抽取 -> 候选先行词生成 -> 规则打分 -> 歧义检测 -> 文本改写 -> Web 展示/数据集评估"
    )
    add_numbered(doc, [
        "用户在 Web 页面输入中文文本。",
        "Mention 抽取模块识别候选实体和代词。",
        "Resolver 模块为每个代词生成候选先行词并计算分数。",
        "系统输出指代关系、候选得分和规则理由。",
        "Rewriter 模块根据消解结果生成改写文本。",
        "评估模块在数据集上计算指标并输出错误分析。",
    ])
    add_para(
        doc,
        "系统设计的核心原则是可解释性。每条指代关系不仅给出最终答案，还展示候选实体得分和规则理由，"
        "便于课堂展示和报告分析。"
    )
    doc.add_heading("4.1 页面功能设计", level=2)
    add_bullets(doc, [
        "单文本分析：输入一句中文文本后展示实体、代词、指代关系和改写结果。",
        "原文高亮：对先行词和代词进行不同颜色标记，帮助观察系统判断。",
        "候选得分表：展示每个候选实体的得分与规则解释。",
        "数据集评估：选择 demo、dev、test、real_corpus、train 等数据集计算指标。",
        "消融实验：关闭不同规则信号，观察 F1 的变化。",
        "数据标注页：支持手动标注文本中的 pronoun 与 antecedent。",
    ])
    doc.add_heading("4.2 系统架构与数据流", level=2)
    add_table(
        doc,
        ["阶段", "输入", "处理模块", "输出"],
        [
            ["1", "用户中文文本", "app.py", "待分析文本"],
            ["2", "待分析文本", "mention_extractor.py", "候选实体、代词、组织指代表达"],
            ["3", "候选实体与代词", "resolver.py", "指代关系、候选得分、规则理由、歧义标记"],
            ["4", "消解结果", "rewriter.py", "显式化后的改写句子"],
            ["5", "数据集与预测结果", "evaluator.py", "Precision、Recall、F1、错误样本"],
            ["6", "结果对象", "Streamlit 页面", "高亮文本、关系表、得分表、实验图表"],
        ],
        [900, 2200, 2400, 3860],
    )

    doc.add_heading("5. 核心算法与规则", level=1)
    add_para(
        doc,
        "本项目当前采用轻量级规则方法作为 baseline。规则方法虽然不如大型预训练模型复杂，但具有实现成本低、"
        "运行速度快、解释性强和适合课堂展示等优点。主要规则包括："
    )
    add_table(
        doc,
        ["规则信号", "含义", "作用"],
        [
            ["距离规则", "候选实体与代词之间的字符距离", "距离越近，通常越可能是先行词。"],
            ["位置规则", "候选实体是否出现在代词之前", "降低后文实体被误选的概率。"],
            ["类型规则", "人称、物体、组织指代表达与实体类型匹配", "处理“他/她/它/该公司”等主要模式。"],
            ["性别规则", "根据已知人名或代词性别判断", "辅助处理“他”和“她”的区分。"],
            ["角色规则", "识别“替/帮/为”“对/向/跟说”等模式", "处理动作角色和部分歧义样本。"],
            ["歧义检测", "最高分与次高分差距较小时标记歧义", "提示人工确认，避免错误自动改写。"],
        ],
        [1700, 3900, 3760],
    )
    doc.add_heading("5.1 打分方式", level=2)
    add_para(
        doc,
        "系统对每个“代词-候选先行词”组合计算一个综合分数。当前版本主要由距离、位置、类型和性别信号构成，"
        "并在特定句式中加入角色规则修正。简化表达如下："
    )
    add_code_block(
        doc,
        "score = distance_score + position_score + type_score + gender_score + role_bonus"
    )
    add_para(
        doc,
        "当最高分与次高分差距较小时，系统会标记为歧义关系。歧义关系仍会展示在页面中，但改写模块会跳过该代词，"
        "避免在不确定情况下生成错误改写。"
    )

    doc.add_heading("6. 数据集构建与训练接入", level=1)
    samples = load_json("data/samples.json")
    dev = load_json("data/dev.json")
    test = load_json("data/test.json")
    real = load_json("data/real_corpus.json")
    train = load_json("data/train.json")
    add_table(
        doc,
        ["数据集", "样本数", "用途"],
        [
            ["demo.json", "10", "页面演示样例。"],
            ["dev.json", str(len(dev)), "规则调试和开发阶段验证。"],
            ["test.json", str(len(test)), "构造测试集，用于最终基础评估。"],
            ["real_corpus.json", str(len(real)), "真实语料风格样本，用于观察泛化问题。"],
            ["samples.json", str(len(samples)), "完整构造样本备份。"],
            ["train.json", str(len(train)), "由新标注集生成，用于规则增强和错误分析。"],
        ],
        [2200, 1500, 5660],
    )
    add_para(
        doc,
        "新标注集 data/new_set.json 共 85 条样本，原始数据未包含 rewrite 字段。项目新增 "
        "scripts/prepare_training_data.py 脚本，自动根据 coreference 标注生成 rewrite，"
        "并合并输出 data/train.json。当前项目中的“训练”并非神经网络参数训练，而是将标注数据中的"
        " pronoun 和 antecedent 并入运行时词表，从而增强实体和指代表达识别能力。"
    )
    doc.add_heading("6.1 数据标注格式", level=2)
    add_para(doc, "每条样本采用 JSON 格式保存，包含原文、指代关系和改写结果。典型格式如下：")
    add_code_block(
        doc,
        '{\n'
        '  "text": "小明把书给了小红，因为她需要它。",\n'
        '  "coreference": [\n'
        '    {"pronoun": "她", "antecedent": "小红"},\n'
        '    {"pronoun": "它", "antecedent": "书"}\n'
        '  ],\n'
        '  "rewrite": "小明把书给了小红，因为小红需要书。"\n'
        '}'
    )

    doc.add_heading("7. 实验结果", level=1)
    doc.add_heading("7.1 新训练集实验", level=2)
    add_table(
        doc,
        ["阶段", "样本数", "真实关系数", "正确关系数", "Precision", "Recall", "F1", "错误样本"],
        [
            ["组织指代增强前", "85", "162", "129", "0.7963", "0.7963", "0.7963", "34"],
            ["组织指代增强后", "85", "162", "142", "0.8765", "0.8765", "0.8765", "21"],
        ],
        [1500, 1000, 1200, 1200, 1100, 1100, 1000, 1260],
    )
    add_para(
        doc,
        "实验结果表明，新增组织指代规则后，训练集 F1 提升 0.0802，错误样本从 34 个减少到 21 个。"
        "这说明“该院”“该校”“这家公司”“这家车企”等组织型指代表达是影响系统表现的重要因素。"
    )
    add_para(
        doc,
        "需要说明的是，train.json 来自新标注集，因此该结果主要反映规则增强对当前真实风格样本的拟合与改进效果。"
        "为了避免过度乐观，报告同时保留 test 和 real_corpus 上的消融实验结果，用于观察规则在不同数据分布上的稳定性。"
    )

    doc.add_heading("7.2 消融实验", level=2)
    add_table(
        doc,
        ["数据集", "实验设置", "正确/真实", "F1", "结论"],
        [
            ["test", "完整规则", "90/90", "1.0000", "构造测试集上完整规则表现稳定。"],
            ["test", "去掉类型规则", "44/90", "0.4889", "类型规则对构造样本最关键。"],
            ["real_corpus", "完整规则", "44/46", "0.9565", "真实语料风格样本仍有少量错误。"],
            ["real_corpus", "去掉距离规则", "32/46", "0.6957", "真实语料中距离规则影响明显。"],
            ["real_corpus", "仅类型+性别", "32/46", "0.6957", "仅靠类型和性别不足以处理复杂上下文。"],
        ],
        [1400, 2100, 1500, 1200, 3160],
    )

    doc.add_heading("8. 错误分析", level=1)
    add_para(
        doc,
        "在训练集上，增强后的系统仍有 21 个错误样本。错误并非集中在单一问题上，而是由多种中文语言现象共同造成。"
    )
    add_table(
        doc,
        ["错误类型", "样本数", "说明"],
        [
            ["标注粒度不一致：它/它们", "6", "部分样本原文为“它们”，gold 标注为“它”，导致评估层面产生漏判和误判。"],
            ["多人物角色/同形代词", "4", "多个“他”或“她”出现时，仅用文本表面形式难以区分位置。"],
            ["组织指代仍需加强", "4", "部分“该院/该平台”仍会偏向最近的展品或产品实体。"],
            ["过度预测", "2", "系统预测了 gold 中未标注的指代关系。"],
            ["身份称谓：老师/弟子", "2", "需要身份知识和常识推理。"],
            ["其他语义判断", "2", "需要更细粒度语义理解。"],
            ["对话角色：我/你", "1", "第一、第二人称依赖说话人和受话人角色。"],
        ],
        [2700, 1100, 5560],
    )
    add_para(
        doc,
        "典型错误包括：“小美会见了小刚，他对她说了句玩笑话”中系统可能交换两个角色；"
        "“孔子问子路，弟子回答了老师的问题”需要识别“弟子”和“老师”的身份关系；"
        "“它们”与 gold 中的“它”不一致则属于标注粒度问题。"
    )
    doc.add_heading("8.1 错误分析结论", level=2)
    add_bullets(doc, [
        "规则方法可以覆盖大量明确指代，但对需要常识和事件角色理解的句子仍有限。",
        "当前评估以 pronoun 文本为主，遇到多个相同代词时容易丢失位置信息。",
        "标注规范会直接影响指标，例如原文为“它们”而 gold 写为“它”时会造成额外错误。",
        "组织型指代表达是本轮提升最大的方向，但仍需要继续补充机构名识别和短语类型判断。",
    ])

    doc.add_heading("9. 可视化展示与部署", level=1)
    add_para(
        doc,
        "项目使用 Streamlit 构建可视化页面。页面支持单文本分析、数据集评估、消融实验和数据标注。"
        "用户输入一句中文后，系统会展示原文高亮、候选实体、代词、指代关系、候选实体得分和改写结果。"
        "对于歧义样本，系统会给出提示并跳过自动改写，从而避免生成误导性结果。"
    )
    add_para(
        doc,
        "项目已经支持云服务器部署。当前简化部署方式为云服务器运行 Streamlit，并通过公网 IP 与端口访问；"
        "正式展示时可进一步使用 Nginx 反向代理和域名绑定，使访问地址更简洁。"
    )
    doc.add_heading("9.1 运行方式", level=2)
    add_para(doc, "本地运行命令如下：")
    add_code_block(doc, "pip install -r requirements.txt\nstreamlit run app.py")
    add_para(doc, "云服务器展示时，可在服务器中拉取 GitHub 仓库并运行：")
    add_code_block(
        doc,
        "git clone https://github.com/gingko123/chinese-coreference-rewrite.git\n"
        "cd chinese-coreference-rewrite\n"
        "python3 -m venv .venv\n"
        "source .venv/bin/activate\n"
        "pip install -r requirements.txt\n"
        "streamlit run app.py --server.address 0.0.0.0 --server.port 8501"
    )
    add_para(
        doc,
        "在课堂展示中，用户打开网页后输入一句中文文本，即可看到候选实体、代词、指代关系、候选得分、歧义提示和改写结果。"
    )
    doc.add_heading("9.2 测试用例与结果说明", level=2)
    add_table(
        doc,
        ["编号", "输入文本", "期望结果", "系统表现"],
        [
            ["1", "马斯克整理了书，他准备再次确认它。", "他->马斯克；它->书", "可正确识别人物与物体指代，并生成“马斯克准备再次确认书”。"],
            ["2", "马云会见了马斯克，他对他说了句玩笑话。", "存在两个“他”的歧义", "系统标记歧义，提示人工确认，避免自动错误改写。"],
            ["3", "该公司发布了新手机，它获得了用户关注。", "该公司/它需区分组织与产品", "组织指代规则增强后，可更稳定处理“该公司”等表达。"],
            ["4", "小明把书给了小红，因为她需要它。", "她->小红；它->书", "基础人物与物体指代可正确改写。"],
            ["5", "博物院展出了瓷器，该院还开放了新的展厅。", "该院->博物院", "组织型指代表达可被识别，是本轮训练增强的重点样例。"],
        ],
        [700, 2850, 2550, 3260],
    )

    doc.add_heading("10. 项目特点", level=1)
    add_bullets(doc, [
        "完整性：覆盖数据、算法、评估、错误分析、可视化和部署。",
        "可解释性：每条关系都有候选得分和规则理由。",
        "可扩展性：预留 HanLP、LTP 和 BERT 语义匹配方向。",
        "可展示性：Streamlit 页面和 PPT 能直观展示系统流程与效果。",
        "真实性：新增真实风格标注集后，结果不再停留在构造样本的满分状态，更能体现实验价值。",
    ])
    doc.add_heading("10.1 与 neuralcoref 类项目的关系", level=2)
    add_para(
        doc,
        "HuggingFace neuralcoref 面向英文神经网络共指消解，本项目则选择中文场景和更轻量的规则增强路线。"
        "两者关注点相似，都是识别文本中的指代关系；但本项目更适合期末作业规模，强调中文数据标注、"
        "规则可解释、改写输出和可视化展示，而不是训练大型端到端模型。"
    )

    doc.add_heading("11. 不足与未来改进", level=1)
    add_table(
        doc,
        ["序号", "不足或问题", "改进方向"],
        [
            ["1", "同形代词难以区分", "引入 span 级评估，记录 pronoun_start 与 antecedent_start。"],
            ["2", "标注粒度存在不一致", "统一按原文完整 span 标注，例如区分“它”和“它们”。"],
            ["3", "复杂语义判断能力有限", "引入 BERT 或 sentence-transformers 计算上下文语义相似度。"],
            ["4", "对话角色识别仍不稳定", "增强“说/问/嘱咐”等事件结构与语义角色规则。"],
            ["5", "数据规模仍然较小", "继续扩展新闻、文学、校园、企业等多场景样本。"],
        ],
        [900, 3300, 5160],
    )

    doc.add_heading("12. 项目分工", level=1)
    add_para(
        doc,
        "本项目按照 50%-50% 的方式进行分工，二人共同完成需求讨论、测试验证和最终材料整理。"
        "具体分工如下表所示，可在提交前将成员姓名替换为真实姓名。"
    )
    add_table(
        doc,
        ["成员", "主要工作", "贡献比例"],
        [
            ["郭英恺", "系统设计、规则消解模块、Streamlit 页面、部署调试。", "50%"],
            ["刘毅恒", "数据标注、训练集整理、实验评估、报告与 PPT 材料整理。", "50%"],
        ],
        [1800, 5760, 1800],
    )

    doc.add_heading("13. 思考与总结", level=1)
    doc.add_heading("13.1 开发问题与解决方案", level=2)
    add_table(
        doc,
        ["问题", "原因", "解决方案"],
        [
            ["HanLP/LTP 环境不熟悉", "中文 NLP 工具安装和模型依赖较重", "先实现规则 baseline，并预留可选后端接口，保证系统可运行。"],
            ["同一句存在多种合理指代", "中文代词依赖上下文与语义角色", "加入歧义检测，分数接近时提示人工确认并跳过自动改写。"],
            ["新增数据缺少 rewrite", "标注集只包含 coreference 关系", "编写 prepare_training_data.py 自动补全 rewrite 并生成 train.json。"],
            ["实验结果需要解释", "单纯指标难以说明问题", "增加消融实验和错误类型归因，展示规则贡献和失败原因。"],
            ["课堂演示需要远程访问", "本地运行不便于展示", "支持云服务器运行 Streamlit，后续可接 Nginx 与域名。"],
        ],
        [2200, 3100, 4060],
    )
    doc.add_heading("13.2 个人收获与反思", level=2)
    add_para(
        doc,
        "通过本项目，可以更直观地体会到 NLP 应用开发并不只是调用模型接口，还包括任务定义、数据格式设计、"
        "算法选择、可视化展示、部署运行和实验解释等完整环节。规则方法虽然能力有限，但在小规模课程项目中具有"
        "可解释、易调试、便于演示的优势；同时，错误分析也说明仅靠表面规则很难解决复杂语义和常识推理问题。"
    )
    add_para(
        doc,
        "后续如果继续完善该项目，应优先从 span 级标注、语义相似度模型和更规范的数据集三个方向推进。"
        "这样既能保留当前系统的可解释展示能力，也能逐步提升复杂场景下的真实语言理解效果。"
    )
    doc.add_heading("13.3 总结", level=2)
    add_para(
        doc,
        "本项目完成了一个中文指代消解与句子改写系统的完整原型。系统以规则方法为基础，"
        "通过 Streamlit 提供可视化展示，并能够在多个数据集上进行评估和错误分析。"
        "在新增 85 条标注数据后，系统自动生成训练集和改写文本，并通过组织指代规则增强将训练集 F1 提升到 0.8765。"
        "虽然当前方法仍难以完全解决多人物同形代词、称谓指代和复杂语义推理问题，但它形成了一个清晰、可解释、"
        "可运行、可展示的中文 NLP 项目闭环，也为后续引入 BERT 语义匹配和 span 级评估奠定了基础。"
    )

    doc.add_heading("参考材料", level=1)
    add_bullets(doc, [
        "项目 README.md 与 IMPLEMENTATION_PLAN.md。",
        "reports/ablation_results.md、reports/training_data_summary.md、reports/training_error_analysis.md。",
        "中文指代消解与句子改写系统_14页汇报PPT.pptx。",
        "ChineseCoreferenceRewriter.pdf。",
    ])

    REPORT_DIR.mkdir(exist_ok=True)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
    print(OUTPUT_PATH)
