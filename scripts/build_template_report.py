from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parents[1]
DATA_DIR = BASE_DIR / "data"
REPORT_DIR = BASE_DIR / "reports"
TEMPLATE_DIR = Path(
    r"C:/Users/30657/Documents/xwechat_files/wxid_843vwkgrw2gj22_03a2/msg/file/2026-06"
)
TEMPLATE_PATH = TEMPLATE_DIR / "期末报告模板.docx"
OUTPUT_PATH = REPORT_DIR / "期末报告_按模板_中文指代消解与句子改写系统.docx"
PIPELINE_IMAGE = Path(
    r"C:/Users/30657/Documents/xwechat_files/wxid_843vwkgrw2gj22_03a2/temp/RWTemp/2026-06/1714780acd71c8319d85d4546345c402/b6ef372a4ad149575cb09ef37f4c59e1.jpg"
)
UI_IMAGE = Path(
    r"C:/Users/30657/AppData/Local/Temp/codex-clipboard-95fce72d-4f8f-43f3-b015-b9edb843f985.png"
)


def set_east_asia_font(run, font_name: str = "宋体") -> None:
    """Set Chinese font for a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def clear_document(doc: Document) -> None:
    """Remove template body content while keeping section/style definitions."""
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(3)

    for name, size in [("Heading 1", 16), ("Heading 2", 15), ("Heading 3", 14)]:
        style = doc.styles[name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_center(doc: Document, text: str, size: int = 16, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-12)
        run = p.add_run("• " + item)
        set_east_asia_font(run)
        run.font.size = Pt(12)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "666666")
        borders.append(tag)
    tbl_pr.append(borders)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "EDEDED")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_east_asia_font(run)
        run.bold = True
        run.font.size = Pt(10.5)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_east_asia_font(run)
            run.font.size = Pt(10.5)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    """Insert an image with a centered caption."""
    if not path.exists():
        add_para(doc, f"（图片未找到：{path}）")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    set_east_asia_font(cap_run)
    cap_run.font.size = Pt(10.5)


def number_headings(doc: Document) -> None:
    """Add template-style numbering to report headings."""
    cn_numbers = "一二三四五六七八九十"
    h1 = 0
    h2 = 0
    h3 = 0
    prefix_pattern = re.compile(r"^(?:[一二三四五六七八九十]+、|（[一二三四五六七八九十]+）|\d+[.．、])\s*")
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        if not style_name.startswith("Heading"):
            continue
        title = prefix_pattern.sub("", paragraph.text.strip())
        if not title:
            continue
        if style_name == "Heading 1":
            h1 += 1
            h2 = 0
            h3 = 0
            prefix = f"{cn_numbers[h1 - 1]}、" if h1 <= len(cn_numbers) else f"{h1}、"
        elif style_name == "Heading 2":
            h2 += 1
            h3 = 0
            prefix = f"（{cn_numbers[h2 - 1]}）" if h2 <= len(cn_numbers) else f"（{h2}）"
        elif style_name == "Heading 3":
            h3 += 1
            prefix = f"{h3}. "
        else:
            continue
        paragraph.text = prefix + title


def load_json(name: str) -> list[dict]:
    path = DATA_DIR / name
    if not path.exists():
        return []
    return json.loads(path.read_text(encoding="utf-8"))


def add_cover(doc: Document) -> None:
    doc.add_paragraph()
    add_center(doc, "华东师范大学", 20, True)
    doc.add_paragraph()
    add_center(doc, "《自然语言处理》", 18, True)
    doc.add_paragraph()
    add_center(doc, "期末报告", 20, True)
    doc.add_paragraph()
    add_center(doc, "标题：中文指代消解与句子改写系统", 15, True)
    doc.add_paragraph()
    add_center(doc, "小组成员：", 13, False)
    add_center(doc, "姓名：郭英恺   学号：10245102433  学院：____________", 12)
    add_center(doc, "姓名：刘毅恒   学号：10245102406  学院：____________", 12)
    doc.add_paragraph()
    add_center(doc, "2026年6月25日", 12)
    doc.add_page_break()


def build_report() -> None:
    doc = Document(str(TEMPLATE_PATH))
    clear_document(doc)
    style_document(doc)

    add_cover(doc)

    samples = load_json("samples.json")
    dev = load_json("dev.json")
    test = load_json("test.json")
    real = load_json("real_corpus.json")
    train = load_json("train.json")
    clue = load_json("clue_wsc.json")
    clue_train = load_json("clue_wsc_train.json")
    clue_dev = load_json("clue_wsc_dev.json")

    doc.add_heading("项目背景", level=1)
    doc.add_heading("问题描述与实际需求", level=2)
    add_para(
        doc,
        "中文文本中大量存在代词、省略和指代表达，例如“他、她、它、该公司、该院、这家车企”等。"
        "如果系统无法判断这些表达对应的先行实体，后续的问答、摘要、信息抽取和文本改写都会受到影响。"
        "本项目围绕中文指代消解任务，进一步加入句子改写功能，使系统不仅能输出“代词指向谁”，还能将原句改写为更清楚的表达。"
    )
    add_para(
        doc,
        "例如输入“小明把书给了小红，因为她需要它”，系统应识别“她→小红”“它→书”，并改写为"
        "“小明把书给了小红，因为小红需要书”。这种显式化表达便于课堂展示，也可以作为问答、摘要等下游任务的预处理步骤。"
    )
    doc.add_heading("项目目标", level=2)
    add_bullets(doc, [
        "实现一个可运行的中文指代消解与句子改写系统。",
        "识别中文句子中的候选实体、代词和组织型指代表达。",
        "通过可解释规则对候选先行词进行打分，并展示候选得分和判断理由。",
        "对可消解代词进行自动改写，对歧义代词进行提示并跳过自动改写。",
        "提供 Web 可视化页面、数据集评估、消融实验和数据标注辅助工具。",
    ])

    doc.add_heading("主要功能", level=1)
    doc.add_heading("系统核心功能", level=2)
    add_table(
        doc,
        ["功能模块", "功能说明"],
        [
            ["单文本分析", "输入一句中文文本，展示候选实体、代词、指代关系、候选得分和改写结果。"],
            ["原文高亮", "对先行词和代词进行不同颜色标记，直观展示系统判断。"],
            ["歧义检测", "当最高分和次高分接近时，提示存在多种可能解释，并跳过自动改写。"],
            ["数据集评估", "批量比较系统预测关系与人工标注关系，计算 Accuracy、Precision、Recall 和 F1。"],
            ["消融实验", "关闭距离、位置、类型、性别等规则，观察指标变化，分析各规则贡献。"],
            ["数据标注器", "通过字符位置选择 pronoun 与 antecedent，导出 JSON 格式标注数据。"],
        ],
    )
    doc.add_heading("功能模块划分与架构图", level=2)
    add_para(doc, "系统采用规则驱动的流水线结构，主要数据流如下：")
    add_para(
        doc,
        "输入中文文本 → Mention 抽取 → 候选先行词生成 → 规则打分 → 歧义检测 → 文本改写 → Web 展示/数据集评估。"
    )
    add_image(doc, PIPELINE_IMAGE, "图1 系统总体流水线与数据标注辅助流程", width=6.4)
    add_table(
        doc,
        ["阶段", "输入", "处理模块", "输出"],
        [
            ["1", "用户中文文本", "app.py", "待分析文本"],
            ["2", "待分析文本", "mention_extractor.py", "候选实体、代词和组织指代表达"],
            ["3", "实体与代词", "resolver.py", "指代关系、候选得分、规则理由"],
            ["4", "消解结果", "rewriter.py", "显式化后的改写句子"],
            ["5", "数据集与预测结果", "evaluator.py", "评价指标与错误样本"],
        ],
    )
    doc.add_heading("使用场景与目标用户", level=2)
    add_bullets(doc, [
        "NLP 课程学习者：通过可视化页面理解中文指代消解流程。",
        "中文文本处理者：将含糊代词替换为明确实体，提高文本可读性。",
        "问答/摘要系统开发者：在下游任务前显式化上下文指代关系。",
        "小规模 NLP 项目开发者：借助标注器、评估页和错误分析快速迭代规则。",
    ])
    doc.add_heading("应用落地设计", level=2)
    add_para(
        doc,
        "本项目可以落地为“中文文本可读性增强工具”。系统并不只输出算法标签，而是把文本中的代词、组织指代和物体指代显式化，"
        "帮助用户理解“他、她、它、其、该公司、该院”等表达具体指向的对象。对于阅读者而言，系统降低了长句和跨句理解成本；"
        "对于下游 NLP 系统而言，改写后的文本语义更明确，适合作为问答、摘要和信息抽取任务的前置处理。"
    )
    add_table(
        doc,
        ["落地场景", "具体作用", "当前实现基础"],
        [
            ["教育阅读辅助", "高亮人物、机构和物体之间的指代关系，帮助学生理解长句和篇章关系。", "单文本分析页、原文高亮、指代关系表。"],
            ["新闻/公告理解", "将“该公司、其、该院”等表达还原为明确实体，提升公告和新闻文本的可读性。", "组织指代词表、候选实体打分、句子改写模块。"],
            ["智能问答预处理", "在进入问答或摘要模型前显式化上下文，减少“他是谁”“该机构指谁”的歧义。", "rewrite 输出、JSON 数据格式、批量评估接口。"],
            ["语料标注辅助", "提供网页标注器，支持构建小规模中文指代关系数据集。", "数据标注页、导出数据、训练/测试集生成脚本。"],
        ],
    )
    add_para(
        doc,
        "因此，本项目的落地点不是直接替代人工判断，而是作为阅读辅助和 NLP 预处理模块使用。系统保留候选得分、规则理由和歧义提示，"
        "使结果可以被人工校验，这也降低了在真实场景中错误改写带来的风险。"
    )
    add_para(
        doc,
        "针对新闻/公告文本，项目进一步加入了长文本处理模式。系统会将输入切分为标题、段落和句子，"
        "把标题中的主要机构作为篇章实体记忆，并在每个句子中结合前文实体窗口进行判断。"
        "这样可以缓解长文本中候选实体过多、距离规则失效和组织指代频繁出现的问题，使“该公司、其、该院”等表达更容易被还原为明确实体。"
    )

    doc.add_heading("技术方案", level=1)
    doc.add_heading("技术选型与理由", level=2)
    add_table(
        doc,
        ["技术/工具", "选择理由"],
        [
            ["Python", "生态成熟，适合快速实现 NLP 规则、数据处理和实验评估。"],
            ["Streamlit", "能够快速构建交互式 Web 页面，适合课堂演示和远程部署。"],
            ["规则驱动 baseline", "实现成本低、可解释性强，便于展示每个判断的依据。"],
            ["JSON 数据集", "结构简单，便于保存 text、coreference、rewrite 等字段。"],
            ["HanLP/LTP 预留接口", "作为后续增强方向，可提供分词、词性标注和命名实体识别能力。"],
            ["云服务器部署", "通过公网网页进行程序演示，便于老师和同学直接访问。"],
        ],
    )
    doc.add_heading("系统架构设计", level=2)
    add_para(
        doc,
        "系统前端由 Streamlit 页面负责，包括单文本分析、数据集评估、消融实验和数据标注四个页面。"
        "后端核心模块包括 mention_extractor、resolver、rewriter 和 evaluator。Mention 抽取模块识别实体和代词；"
        "resolver 模块对每个候选先行词打分；rewriter 模块根据消解结果进行替换；evaluator 模块负责批量评估和错误分析。"
    )
    doc.add_heading("部署说明", level=2)
    add_para(doc, "本地运行方式如下：")
    add_para(doc, "pip install -r requirements.txt；streamlit run app.py。")
    add_para(
        doc,
        "云服务器部署时，拉取 GitHub 仓库后创建虚拟环境，安装依赖，并运行："
        "streamlit run app.py --server.address 0.0.0.0 --server.port 8501。"
        "当前项目也可通过 systemd 配置为常驻服务，使服务器重启后自动恢复网页访问。"
    )
    doc.add_heading("多端口部署实现与效果对比", level=2)
    add_para(
        doc,
        "为了适应课堂演示、远程访问和长期运行三种需求，项目设计了多种访问端口和部署方式。"
        "本地开发阶段可以使用 Streamlit 默认端口 8501；云服务器演示阶段可以开放 8501 端口进行公网访问；"
        "正式展示阶段可以通过 Nginx 将 80/443 端口反向代理到 8501，使访问形式更接近普通网页。"
    )
    add_table(
        doc,
        ["部署方式", "访问端口", "访问形式", "优点", "不足"],
        [
            ["本地 Streamlit", "8501", "http://localhost:8501", "启动简单，适合开发调试。", "只能本机访问，不适合课堂远程展示。"],
            ["云服务器直连", "8501", "http://公网IP:8501", "配置少，部署速度快，适合临时演示。", "需要安全组放行端口，网址形式不够正式。"],
            ["systemd 常驻服务", "8501", "http://公网IP:8501", "SSH 断开后仍可运行，服务器重启后可自动恢复。", "仍然暴露应用端口，缺少域名和 HTTPS。"],
            ["Nginx 反向代理", "80/443", "http(s)://域名", "访问方式正式，可接入 HTTPS，适合最终展示。", "配置复杂度更高，需要域名和证书。"],
        ],
    )
    add_para(
        doc,
        "从效果上看，8501 端口直连能够最快完成演示验证；systemd 解决了“关闭终端后网页失效”的问题；"
        "Nginx + 域名 + HTTPS 则更适合作为正式部署方案。对于本课程项目，当前已经完成云服务器直连和 systemd 常驻运行，"
        "后续若继续完善，可进一步接入 Nginx 和 HTTPS，提高访问稳定性与展示规范性。"
    )

    doc.add_heading("关键算法/模型原理说明", level=1)
    doc.add_heading("Mention 抽取机制", level=2)
    add_para(
        doc,
        "本项目没有强依赖复杂分词模型，而是采用“规则词表 + 字符级 span 匹配 + 可选 NLP 后端”的方式进行 Mention 抽取。"
        "系统会记录每个 mention 的文本、类型、起止位置，方便后续高亮展示和句子改写。"
    )
    add_para(
        doc,
        "例如“小明把书给了小红，因为她需要它”会抽取出 PERSON:小明、OBJECT:书、PERSON:小红、PRONOUN:她、PRONOUN:它。"
    )
    doc.add_heading("规则打分与消解", level=2)
    add_para(
        doc,
        "系统为每个“代词-候选先行词”组合计算综合分数。分数主要由距离、位置、类型、性别和角色模式组成。"
        "类型匹配和距离规则是核心信号，例如“她”更倾向人物，“它”更倾向物体，“该院/该公司”更倾向组织。"
    )
    add_para(doc, "简化表达为：score = distance_score + position_score + type_score + gender_score + role_bonus。")
    doc.add_heading("创新点与技术贡献", level=2)
    add_para(
        doc,
        "本项目的创新性不体现在训练一个大型端到端神经网络模型，而体现在将中文指代消解任务设计成一个可解释、"
        "可改写、可评估、可部署的完整 NLP 应用系统。传统的共指消解示例通常只输出“代词—先行词”标签，"
        "而本项目进一步将消解结果转化为自然语言改写结果，使用户能够直接看到系统理解后的句子表达。"
        "例如“马斯克整理了书，他准备再次确认它”可以被改写为“马斯克整理了书，马斯克准备再次确认书”，"
        "这种输出比单纯显示标签更适合课堂演示和下游文本处理。"
    )
    add_para(
        doc,
        "第二个创新点是可解释规则打分机制。系统不是直接给出黑盒预测，而是保留每个候选实体的分数和规则理由，"
        "包括距离、位置、类型、性别和角色模式等信号。这样可以回答“系统为什么选择这个先行词”的问题，"
        "也便于开发者根据错误样本继续调整规则。对于课程项目而言，可解释性能够帮助展示算法过程，而不只是展示最终结果。"
    )
    add_para(
        doc,
        "第三个创新点是歧义检测和安全改写机制。中文中经常出现多个同形代词，例如“马云会见了马斯克，他对他说了句玩笑话”，"
        "两个“他”的指向都可能存在不同解释。系统在最高分和次高分接近时会标记为歧义，并提示人工确认，"
        "同时跳过自动改写，避免生成误导性文本。这体现了系统在不确定场景下的保守处理策略。"
    )
    add_para(
        doc,
        "第四个创新点是数据闭环设计。项目不仅包含单句分析 Demo，还提供数据标注器、数据集评估、消融实验和错误分析。"
        "标注器可以通过字符级 span 选择生成 pronoun 与 antecedent 的对应关系；评估模块可以批量计算 Accuracy、"
        "Precision、Recall 和 F1；消融实验可以验证距离、位置、类型、性别等规则的实际贡献。"
        "这些模块共同构成从数据构建到规则优化再到结果展示的迭代流程。"
    )
    add_para(
        doc,
        "第五个创新点是面向中文场景的组织指代增强。针对“该院、该校、该平台、这家公司、这家车企”等中文常见表达，"
        "项目扩展了组织型指代词表和实体类型判断。新增真实风格标注集后，通过组织指代规则增强，训练集 F1 从 0.7963 "
        "提升到 0.8765，说明该改进对中文实际语料具有明确贡献。"
    )
    add_para(
        doc,
        "总体来看，本项目的技术贡献在于提供了一个轻量级、可解释、可运行的中文指代消解 baseline，并将其与句子改写、"
        "可视化展示、数据标注和云端部署结合起来。后续若引入 BERT 语义相似度、依存句法或学习排序模型，"
        "可以在当前可解释框架上继续提升复杂语义场景下的表现。"
    )
    doc.add_page_break()
    doc.add_heading("数据集与实验结果", level=2)
    add_table(
        doc,
        ["数据集", "样本数", "用途"],
        [
            ["demo.json", "10", "页面演示样例"],
            ["dev.json", str(len(dev)), "规则调试和开发验证"],
            ["test.json", str(len(test)), "基础评估"],
            ["real_corpus.json", str(len(real)), "真实语料风格泛化观察"],
            ["samples.json", str(len(samples)), "完整构造样本"],
            ["train.json", str(len(train)), "新标注集生成，用于规则增强和错误分析"],
            ["clue_wsc.json", str(len(clue)), "CLUEWSC2020 公开集转换样本，用于公开数据集实验"],
            ["clue_wsc_train/dev", f"{len(clue_train)}/{len(clue_dev)}", "公开集训练/验证切分，用于区分训练拟合与泛化效果"],
        ],
    )
    add_table(
        doc,
        ["实验设置", "正确/真实", "F1", "结论"],
        [
            ["训练集增强前", "129/162", "0.7963", "组织指代规则不足，错误较多。"],
            ["训练集增强后", "142/162", "0.8765", "组织指代表达增强后效果提升。"],
            ["real_corpus 完整规则", "44/46", "0.9565", "真实语料风格样本表现较稳定。"],
            ["real_corpus 去掉距离规则", "32/46", "0.6957", "距离规则对系统影响明显。"],
        ],
    )
    doc.add_heading("公开集与最新补充实验", level=3)
    add_para(
        doc,
        "为增强实验可信度，项目进一步接入 CLUEWSC2020 公开中文指代消解数据集。该数据集更偏 Winograd Schema，"
        "不少样本需要常识语义推理，因此比项目自建样本更能暴露规则系统的局限。当前系统将 CLUEWSC2020 转换为项目统一的 "
        "text/coreference/rewrite 格式，并固定切分为训练集和验证集。"
    )
    add_table(
        doc,
        ["实验设置", "Gold", "Pred", "Correct", "Precision", "Recall", "F1", "说明"],
        [
            ["CLUE dev 接入前", "96", "125", "7", "0.0560", "0.0729", "0.0633", "公开集验证切分上的原始规则表现。"],
            ["CLUE dev 仅扩词表后", "96", "151", "5", "0.0331", "0.0521", "0.0405", "候选覆盖增加，但常识推理能力不足。"],
            ["CLUE 全量仅扩词表后", "449", "726", "263", "0.3623", "0.5857", "0.4477", "公开集整体规则 baseline。"],
            ["CLUE 全量记忆增强后", "449", "504", "358", "0.7103", "0.7973", "0.7513", "训练记忆增强后的拟合效果，应与泛化结果分开说明。"],
            ["项目 test 接入后", "90", "90", "80", "0.8889", "0.8889", "0.8889", "项目自建测试集表现稳定。"],
            ["real_corpus 接入后", "46", "46", "34", "0.7391", "0.7391", "0.7391", "真实风格语料难度更高。"],
        ],
    )
    add_para(
        doc,
        "该实验说明：公开集接入能够提升系统的数据来源和实验说服力，但仅靠规则扩词表无法充分解决 CLUEWSC 中的常识推理问题。"
        "训练记忆增强可以显著提高全量公开集上的指标，但它本质上属于 supervised exact-match fitting，适合用于展示训练集拟合能力，"
        "不能等同于开放域泛化能力。后续如果引入 BERT 或 sentence-transformers 语义匹配模型，可以在当前规则候选基础上继续提高复杂样本表现。"
    )
    doc.add_heading("基本模型与 HanLP/LTP 对比", level=2)
    add_para(
        doc,
        "为了让项目从单一规则系统扩展为可比较的实验平台，系统增加了 backend 机制。"
        "规则版 baseline 作为基本模型始终可运行；HanLP 和 LTP/legacy 作为可选增强后端，"
        "在安装并配置模型路径后可以参与 Mention 抽取，并与 baseline 使用同一套 resolver、rewriter 和 evaluator 进行横向比较。"
    )
    add_para(
        doc,
        "本项目中的“模型对比”不是简单比较三个独立系统，而是固定后半段指代消解流程，只替换前端 Mention 抽取层。"
        "也就是说，baseline、HanLP 和 LTP 都会进入同一套候选生成、规则打分、歧义检测和改写模块。"
        "这样设计的好处是可以更清楚地观察中文分词、命名实体识别和词性分析工具对最终指代消解结果的影响，"
        "避免把差异混淆为后续打分规则的变化。"
    )
    add_table(
        doc,
        ["后端", "实现方式", "当前效果展示", "适合说明的问题"],
        [
            ["规则版 baseline", "内置词表、字符 span 匹配、规则打分。", "无需额外模型，当前 test F1 为 0.9667，real_corpus F1 为 0.9565。", "证明系统基本功能可运行，并提供可解释判断过程。"],
            ["HanLP 增强接口", "通过 HANLP_MODEL 配置模型，抽取命名实体后与规则实体合并；未安装时使用 HanLP-style fallback。", "当前环境使用启发式 fallback，偏机构名、品牌名、人名抽取。", "观察 NER 风格实体抽取对结果的影响。"],
            ["LTP/legacy 接口", "支持 LTP_MODEL，也兼容 pyltp 的 cws/pos/ner legacy 模型路径；未安装时使用 LTP-style fallback。", "当前环境使用启发式 fallback，偏分词短语和名词块抽取。", "展示分词/名词块策略对 Mention 抽取的影响。"],
        ],
    )
    add_table(
        doc,
        ["比较维度", "规则版 baseline", "HanLP", "LTP/legacy"],
        [
            ["实体发现方式", "依赖项目内置词表、姓氏规则、组织后缀和物体词表。", "依赖预训练中文多任务模型输出的人名、机构名、地名等实体。", "依赖分词、词性标注和命名实体识别结果，legacy 版本使用 cws/pos/ner 三个模型。"],
            ["优点", "轻量、稳定、可解释，离线即可运行。", "开放领域实体召回能力更强，对真实新闻和公告文本更友好。", "传统中文 NLP 工具链成熟，便于展示分词、词性和 NER 对任务的影响。"],
            ["不足", "词表外实体容易漏识别，复杂短语边界不稳定。", "模型依赖较重，首次下载和运行成本较高，输出格式需要适配。", "安装和模型路径配置较繁琐，Windows 环境下 legacy 依赖可能不稳定。"],
            ["对本任务的预期影响", "适合课程 baseline 和错误分析。", "可能提高人物、机构、地点实体的召回率，从而提升 Recall。", "可能改善实体边界和组织名识别，但效果取决于模型版本和语料领域。"],
        ],
    )
    add_para(
        doc,
        "对比实验采用相同数据集和相同评价指标。系统首先读取标注数据中的 gold coreference，"
        "然后分别使用 rule、hanlp、ltp 三种 backend 预测代词—先行词关系，最后统计 Precision、Recall 和 F1。"
        "其中 Precision 反映预测关系中有多少是正确的，Recall 反映真实指代关系中有多少被找回，F1 则综合衡量准确性和召回能力。"
    )
    add_table(
        doc,
        ["实验条件", "设置说明"],
        [
            ["固定部分", "resolver.py、rewriter.py、evaluator.py 保持一致，只改变 Mention 抽取后端。"],
            ["数据集", "dev.json、test.json、real_corpus.json、samples.json 和 train.json。"],
            ["评价指标", "Precision、Recall、F1、正确/真实关系数、错误样本数。"],
            ["回退机制", "未安装 HanLP/LTP 或未配置模型路径时，系统使用同名启发式 fallback，避免程序无法运行，同时保留可对比差异。"],
        ],
    )
    add_table(
        doc,
        ["数据集", "rule baseline F1", "HanLP 状态", "LTP/legacy 状态", "结果解读"],
        [
            ["CLUEWSC2020", "0.4477", "HanLP-style fallback F1=0.4405", "LTP-style fallback F1=0.4293", "关闭训练记忆后比较 Mention 抽取策略；CLUEWSC 偏常识推理，NER 增强不一定提升 F1。"],
            ["CLUEWSC2020 + 记忆增强", "0.7513", "0.7513", "0.7513", "训练记忆会抹平后端差异，因此模型对比页默认关闭该选项。"],
            ["test", "0.8889", "可在页面实时运行", "可在页面实时运行", "项目自建测试集用于验证系统基本功能。"],
            ["real_corpus", "0.7391", "可在页面实时运行", "可在页面实时运行", "真实语料风格样本更能暴露开放域错误。"],
        ],
    )
    add_para(
        doc,
        "新增的模型对比页会展示每个后端的可用状态、Precision、Recall、F1、相对 baseline 的 ΔF1 和说明信息。"
        "为了避免训练记忆增强抹平不同后端之间的差异，模型对比页默认关闭训练记忆，只比较不同 Mention 抽取策略对最终结果的影响。"
    )
    add_para(
        doc,
        "现阶段的实验结果应表述为 rule baseline、HanLP-style fallback 与 LTP-style fallback 的对比。"
        "由于提交环境尚未安装真实 HanLP/LTP 模型，系统不会声称已经获得真实模型提升，而是通过两个同名启发式后端展示不同实体抽取策略带来的差异。"
        "在 CLUEWSC2020 上，rule baseline F1 为 0.4477，HanLP-style fallback F1 为 0.4405，LTP-style fallback F1 为 0.4293。"
        "该结果说明：HanLP/LTP 类工具主要增强分词和实体识别，而 CLUEWSC 中很多样本需要常识语义推理，因此实体抽取增强并不必然带来更高 F1。"
        "后续只需要安装模型并设置 HANLP_MODEL、LTP_MODEL 或 PYLTP_CWS_MODEL 等路径，即可重新运行脚本生成真实模型对比表。"
    )

    doc.add_heading("界面展示", level=1)
    doc.add_heading("系统界面设计理念", level=2)
    add_para(
        doc,
        "界面设计以课堂演示和可解释性为核心。用户输入中文句子后，页面不会只输出一个结论，而是同时展示原文高亮、"
        "识别结果、指代关系、候选实体得分、歧义提示和改写结果。这样可以让观察者看到系统从输入到输出的完整过程。"
    )
    add_image(doc, UI_IMAGE, "图2 系统单文本分析界面", width=6.3)
    doc.add_heading("主要页面说明", level=2)
    add_table(
        doc,
        ["页面", "展示内容"],
        [
            ["单文本分析", "输入文本、原文高亮、候选实体、代词、指代关系、改写结果。"],
            ["新闻/公告长文本", "标题实体记忆、前文实体窗口、分句处理结果和保守改写结果。"],
            ["数据集评估", "Accuracy、Precision、Recall、F1、错误样本、漏判和误判。"],
            ["消融实验", "完整规则与关闭某类规则后的指标对比，用于分析规则贡献。"],
            ["数据标注", "字符编号、span 选择、添加标注、绑定文本和导出 JSON。"],
        ],
    )
    doc.add_heading("测试用例与结果说明", level=2)
    add_table(
        doc,
        ["编号", "输入文本", "期望结果", "系统表现"],
        [
            ["1", "马斯克整理了书，他准备再次确认它。", "他→马斯克；它→书", "可正确识别人物与物体指代。"],
            ["2", "小明把书给了小红，因为她需要它。", "她→小红；它→书", "基础人物和物体指代可正确改写。"],
            ["3", "马云会见了马斯克，他对他说了句玩笑话。", "存在两个“他”的歧义", "系统提示歧义并跳过自动改写。"],
            ["4", "博物院展出了瓷器，该院还开放了新的展厅。", "该院→博物院", "组织型指代表达可被识别。"],
            ["5", "孔子问子路，弟子回答了老师的问题。", "需要身份关系判断", "暴露称谓和常识知识不足。"],
        ],
    )

    doc.add_heading("思考与总结", level=1)
    doc.add_heading("开发过程中遇到的问题与解决方案", level=2)
    add_para(
        doc,
        "本项目从最初的想法到最终形成可运行的 Web 应用，经历了任务定义、规则设计、数据构建、评估分析和云端部署等多个阶段。"
        "在开发过程中，我们发现 NLP 应用并不是只要实现一个算法函数即可，真正影响系统质量的因素还包括数据标注是否规范、"
        "页面展示是否清楚、错误样本能否被解释，以及程序能否在不同环境中稳定运行。"
    )
    add_table(
        doc,
        ["问题", "原因", "解决方案"],
        [
            ["HanLP/LTP 环境较重", "模型依赖和安装成本较高", "先实现稳定规则版 baseline，并预留可选后端接口。"],
            ["多人物代词存在歧义", "同形代词依赖语义角色", "加入歧义检测，提示人工确认并跳过自动改写。"],
            ["新增数据缺少 rewrite", "标注集只包含指代关系", "编写脚本自动根据 coreference 生成 rewrite。"],
            ["组织指代容易误判", "“该院/该平台”容易指向最近实体", "扩展组织指代表达和组织实体词表。"],
            ["演示需要远程访问", "本地运行不方便课堂展示", "使用云服务器部署 Streamlit，并配置 systemd 常驻服务。"],
        ],
    )
    add_para(
        doc,
        "第一个较明显的问题是中文 NLP 工具链的环境复杂度。HanLP、LTP 等工具能够提供更强的分词、词性标注和命名实体识别能力，"
        "但它们的模型文件、依赖版本和运行环境都比较重。如果一开始就强依赖这些工具，课堂演示时可能因为环境问题导致系统无法运行。"
        "因此本项目选择先实现规则版 baseline，把 HanLP/LTP 作为可选增强接口。这种方案牺牲了一部分泛化能力，但换来了稳定性和可解释性。"
    )
    add_para(
        doc,
        "第二个问题是中文指代关系本身具有歧义。例如“马云会见了马斯克，他对他说了句玩笑话”中，两个“他”都可能对应不同人物。"
        "如果系统强行选择一个先行词并自动改写，就可能产生错误甚至误导性的结果。为了解决这个问题，我们加入了歧义检测机制："
        "当最高分和次高分接近时，系统将该关系标记为歧义，提示用户人工确认，并在改写阶段跳过该代词。"
        "这个设计体现了系统在不确定场景下的保守策略，也提高了展示时的可信度。"
    )
    add_para(
        doc,
        "第三个问题来自数据集构建。新增的标注数据最初只包含 pronoun 和 antecedent 关系，并没有 rewrite 字段。"
        "如果手工为每条样本补写改写结果，工作量较大，也容易出现格式不一致。为此，我们编写了 prepare_training_data.py，"
        "根据 coreference 标注自动生成 rewrite，并合并输出 train.json。这样既提高了数据处理效率，也让后续评估和错误分析更加统一。"
    )
    add_para(
        doc,
        "第四个问题是组织型指代的误判。早期系统在处理“该院、该平台、这家公司、这家车企”等表达时，"
        "容易因为距离规则而选择最近的普通实体，而不是对应的组织实体。通过错误分析，我们发现组织型指代是影响真实风格样本表现的重要因素。"
        "因此后续扩展了组织指代词表，并在 mention 抽取和类型匹配中强化 ORG 类别，使训练集 F1 从 0.7963 提升到 0.8765。"
    )
    add_para(
        doc,
        "第五个问题是部署与展示。Streamlit 在本地运行很方便，但如果只依赖本地电脑，课堂展示时不便于老师和同学访问。"
        "我们将项目部署到云服务器，并通过公网 IP 加端口访问。进一步地，为了解决 SSH 断开或服务器重启后网页不可访问的问题，"
        "我们使用 systemd 配置了常驻服务，使 Streamlit 应用能够随服务器启动自动恢复。"
    )

    doc.add_heading("项目局限性分析", level=2)
    add_para(
        doc,
        "虽然系统已经形成了完整的应用闭环，但当前版本仍然存在明显局限。首先，系统主要依赖规则和词表，"
        "对于训练集和词表中没有出现过的新实体、新组织名或复杂短语，识别能力仍然有限。规则方法适合解释和展示，"
        "但面对开放域文本时，其覆盖范围不如预训练语言模型。"
    )
    add_para(
        doc,
        "其次，当前规则主要利用距离、位置、类型、性别等表层信号，对深层语义和常识推理支持不足。"
        "例如“老师、弟子、领导、下属”等身份称谓关系，需要背景知识才能判断；“他问他”“她提醒她”等对话场景，"
        "需要理解事件角色和说话人关系。仅靠字符距离和类型匹配，很难稳定解决这些问题。"
    )
    add_para(
        doc,
        "再次，当前评估仍以 pronoun 文本和 antecedent 文本为主要比较单位。若一句话中出现多个相同代词，"
        "例如连续出现两个“他”，只比较文本会丢失具体位置差异。更严格的共指消解评估应该记录 pronoun_start、"
        "pronoun_end、antecedent_start 和 antecedent_end，以 span 级别判断系统是否真正找到了正确位置。"
    )
    add_para(
        doc,
        "最后，数据集规模仍然偏小。当前数据能够支持课程项目演示和规则验证，但还不足以代表真实中文文本的全部复杂性。"
        "新闻、校园、企业公告、文学文本、对话文本等不同场景中的指代模式差异较大，后续仍需要继续扩展数据集。"
    )

    doc.add_heading("未来改进方向", level=2)
    add_bullets(doc, [
        "引入 BERT 或 sentence-transformers，将规则候选先行词与上下文语义相似度结合，减少仅靠距离造成的误判。",
        "接入 HanLP 或 LTP 的命名实体识别、词性标注和依存句法分析，提高 Mention 抽取与事件角色判断能力。",
        "改进标注格式，加入 pronoun_start、pronoun_end、antecedent_start、antecedent_end，实现 span 级评估。",
        "扩展真实语料数据集，覆盖新闻、校园、企业、文学、对话等不同文本类型。",
        "优化数据标注器交互，支持鼠标选中文本、删除单条标注、上传原始文本和标注格式校验。",
        "进一步优化改写自然度，避免机械替换导致句子重复或表达不自然。",
        "在部署层面加入 Nginx、域名和 HTTPS，使系统更接近正式 Web 应用。",
    ])
    add_para(
        doc,
        "其中，最优先的改进方向是 span 级评估和语义模型增强。span 级评估可以解决多个同形代词难以区分的问题，"
        "语义模型则可以为规则系统提供更强的上下文理解能力。未来可以将当前规则分数作为特征，再加入语义相似度分数，"
        "形成一个可解释的混合模型。这样既能保留 baseline 的透明性，又能提升复杂语义场景下的准确率。"
    )
    doc.add_heading("个人收获与反思", level=2)
    add_para(
        doc,
        "通过本项目，我们对 NLP 应用开发有了更完整的认识。过去容易把 NLP 项目理解为“调用模型并输出结果”，"
        "但在实际开发中，模型或规则只是其中一环。一个可提交、可展示的应用还需要明确任务边界、设计数据格式、"
        "实现交互界面、构建评估指标、整理错误分析，并考虑如何部署到真实环境中。"
    )
    add_para(
        doc,
        "我们也认识到，可解释性和准确率之间经常需要平衡。端到端模型可能具有更强的泛化能力，但不容易解释每一步决策；"
        "规则系统准确率上限有限，却能清楚展示判断依据。对于本课程项目而言，规则 baseline 更适合展示系统流程、错误来源和改进方向。"
        "后续若要继续提升，可以在现有可解释框架上逐步引入模型，而不是完全替换掉规则结构。"
    )
    add_para(
        doc,
        "此外，数据质量对 NLP 系统影响非常明显。标注粒度不统一、代词位置不明确、rewrite 缺失等问题，都会影响最终指标。"
        "因此，构建数据集并不是简单收集句子，而是需要设计规范、保持一致性，并通过错误分析不断修正。"
        "本项目中的数据标注器虽然比较简易，但它让我们体验到了从原始文本到可评估样本的完整过程。"
    )
    add_para(
        doc,
        "最后，云服务器部署也让我们理解到应用系统和实验脚本的区别。一个脚本在本地能运行，并不代表它适合演示和交付；"
        "只有当系统能够稳定通过网页访问、在服务器重启后自动恢复，并且能被他人直接使用时，它才更接近一个完整应用。"
        "这也是本项目从算法练习走向 NLP 应用开发的重要一步。"
    )

    doc.add_heading("小组分工与贡献", level=1)
    add_table(
        doc,
        ["成员", "主要工作", "贡献比例"],
        [
            ["郭英恺", "系统设计、规则消解模块、Streamlit 页面、部署调试、演示准备。", "50%"],
            ["刘毅恒", "数据标注、训练集整理、实验评估、报告与 PPT 材料整理。", "50%"],
        ],
    )

    REPORT_DIR.mkdir(exist_ok=True)
    number_headings(doc)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
